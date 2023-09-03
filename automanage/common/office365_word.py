#!/usr/bin/env python3
# -*- coding-utf-8 -*-

import os
import glob
import fire
from loguru import logger

from docx import Document
from docx2pdf import convert

class WordFile(object):
    def __init__(self, docx_path = None, pdf_path = None):
        self.docx_path = docx_path
        self.pdf_path = pdf_path
        self.doc = Document(self.docx_path)

    def docx2pdf(self, docx_path, pdf_path = None):
        """
        Convert docx to pdf.

        Args:
            docx_path(str): docx file path
            pdf_path(str): pdf file path.
        Returns:
            None
        """
        logger.info(f'Start converting {docx_path} to {pdf_path}.')
        convert(docx_path, pdf_path)
        logger.info(f'Done converting {docx_path} to {pdf_path}.')

    def docx_info(self, docx_path):
        """
        Get docx file info.
        Args:
            docx_path(str): docx file path
        Returns:
            paragraphs(dict): paragraphs in docx file
            tables(list): tables in docx file
        """
        doc = Document(docx_path)
        # doc = self.doc
        paragraphs = {para.text: para.style.name for para in doc.paragraphs}
        logger.info(f'Showing paragraphs in {docx_path}.')
        docx_file = {}
        for key, value in paragraphs.items():
            logger.debug(f'{value}: {key}')
            # docx_file[key] = value
        logger.info(f'Done showing paragraphs in {docx_path}.')
        logger.debug(set(paragraphs.values()))

        for table in doc.tables:
            # logger.info(table.style.name)
            # # rows = [[cell.text for cell in row.cells] for row in table.rows]
            # # logger.debug('\n'.join(['\t'.join(row) for row in rows]))
            pass

        return paragraphs, doc.tables

    # {'Normal', 'Heading 2', 'Heading 3', 'Heading 1', 'Footer'}
    def modify_text(self, docx_path, keyword, new_keyword):
        doc = Document(docx_path)
        sign = False
        logger.info(f'Start modifying {docx_path} text from {keyword} to {new_keyword}.')
        for para in doc.paragraphs:
            if keyword not in para.text:
                sign = True
                continue
            inlines = para.runs
            new_inlines = [inline.text.replace(keyword, new_keyword) if keyword in inline.text else inline.text for inline in inlines]
            for i, inline in enumerate(inlines):
                inline.text = new_inlines[i]
            if keyword in inline.text:
                logger.debug(f'Found the {keyword} in {docx_path} at {para.style.name}, {para.text}.')
                # logger.debug(f'The runs: {inlines}')
        if sign:
            logger.warning(f'Did not found {keyword} in {docx_path}.')
        else:
            logger.info(f'Done modifying {docx_path} from {keyword} to {new_keyword}.')
        doc.save(docx_path)

    def modify_table(self, docx_path, keyword, new_keyword):
        doc = Document(docx_path)
        logger.info(f'Start modifying {docx_path} table from {keyword} to {new_keyword}.')
        for table in doc.tables:
            if keyword not in table._cells:
                continue
            cells = table.cells
            # cells = [cell.text.replace(keyword, new_keyword) if keyword in cell.text else cell.text for cell in cells]
            logger.debug(f'The cells: {cells}')
            # table.cells = cells
            # if keyword in cell.text:
            #     logger.info(f'Found the {keyword} in {docx_path} at {table.style.name}, {table.cells}.')
            #     # logger.info(f'The runs: {table.cells}')

            # # # for row in table.rows:
            # # #     for cell in row.cells:
            # # #         if keyword not in cell.text:
            # # #             continue
            # # #         cell.text = cell.text.replace(keyword, new_keyword)
            # # #         logger.debug(f'Found the {keyword} in {docx_path} at {table.style.name}, {cell.text}.')
        logger.info(f'Done modifying {docx_path} from {keyword} to {new_keyword}.')
        doc.save(docx_path)
        
    def batch_modify_text(self, dir_path, keyword, new_keyword):
        """
        Modify a file in a batch.
        """
        # files = os.listdir(dir_path)
        # logger.debug(f'Files in batch: {files}')
        # files.sort()
        # for file in files:
        #     path = os.path.join(dir_path, file)
        #     self.modify_text(path, keyword, new_keyword)

        for file_path in glob.glob(os.path.join(dir_path, '*.docx')) + glob.glob(os.path.join(dir_path, '*.doc')):
            try:
                self.modify_text(file_path, keyword, new_keyword)
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")

    def main(self):
        docx_path = './word/20220521.docx'
        pdf_path = './word/20220521.pdf'
        # self.docx2pdf(docx_path, pdf_path)
        # self.docx_info(docx_path)
        keyword = '曾经'
        new_keyword = '评估'
        self.modify_text(docx_path, keyword, new_keyword)
        keyword = '描述'
        self.modify_table(docx_path, keyword, new_keyword)

if __name__ == '__main__':
    fire.Fire(WordFile)
